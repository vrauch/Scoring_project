# Required imports
import nltk
from transformers import AutoTokenizer, AutoModel
from analysis_modules import (
    get_maturity_score,
    to_sentence_case,
    ia_analysis,
    get_embedding,
    compute_cosine_similarity,
    clean_and_normalize_text,
    load_from_db_project,
    save_analysis_result
)
import os
import contextlib
from docx import Document
from docx.shared import Pt  # Import Pt for point size

# model_name = "sentence-transformers/all-MiniLM-L6-v2"
model_name = "gpt2"
tokenizer = AutoTokenizer.from_pretrained(model_name)
model = AutoModel.from_pretrained(model_name)

# Disable parallelism in tokenizers to avoid the warning
os.environ["TOKENIZERS_PARALLELISM"] = "false"

# Load Data
# Prompt the user for the project ID
project_id = input("Please enter the project ID to load data from: ")

# Convert project_id to an integer (if numeric)
try:
    project_id = int(project_id)
except ValueError:
    print("Invalid input. Please enter a numeric project ID.")
    exit(1)  # Exit if input is invalid

# Pass the project_id to the load_from_db function. This will develop an analysis printout for project assessment
rows = load_from_db_project(project_id)

if rows:
    # Setup Output Document (commented out for now)
    # doc = Document()
    # doc.add_heading('Analysis Results', level=1)

    print(f"Data loaded for project ID {project_id}:")

    # Loop through each row (similar to iterating over a DataFrame)
    for row in rows:
        domain = row['Domain']
        capability = row['Capability']
        level = row['Level']
        assessment = row['Assessment']
        criteria = row['Criteria1']
        criteria2 = row['Criteria2']
        project_name = row['project_name']
        capability_id = row['capability_id']  # Assuming capability_id is included in the loaded data

        model_name = "sentence-transformers/all-MiniLM-L6-v2"
        tokenizer = AutoTokenizer.from_pretrained(model_name)
        model = AutoModel.from_pretrained(model_name)

        # Construct the alignment prompt
        prompt = f"""
                    Analyze the following text for alignment with the specified criteria.
                    Text: {assessment}
                    Criteria: {criteria}

                    1. Determine whether the alignment is weak, moderate, or strong.
                    2. Provide a brief, fact-based synopsis (less than 75 words) justifying the alignment rating. Focus on specific points of alignment or gaps relative to the criteria.

                    Output Requirements:
                    - Return only the alignment rating and synopsis in plain text without any labels, headers, or introductory phrases.
                    - Do not use any special formatting, characters, symbols, Markdown, or any form of labeling (e.g., "Alignment:" or "Synopsis:").
                    - Avoid any asterisks (*) or stylistic indicators.
                """

        alignment = ia_analysis(criteria, assessment, prompt)
        alignment = clean_and_normalize_text(alignment)
        alignment = to_sentence_case(alignment)

        maturity_score = get_maturity_score(alignment)
        print(maturity_score)

        # Calculate Semantic Similarity Score
        similarity = compute_cosine_similarity(criteria, assessment, tokenizer, model)
        similarity = round(similarity, 2)



        # Construct the recommendations prompt
        prompt = f"""
                    Analyze the following scenario regarding the progression from Level 1 to Level 2 maturity within a 
                    business context.
                    Assessment at Level 1: {assessment}
                    Criteria for Level 2 Maturity: {criteria2}

                    1. Based on the Level 1 assessment and Level 2 maturity criteria, write a short, fact-based 
                    recommendation synopsis (no more than 75 words) that supports this progression. 
                    Focus on actionable insights and key steps needed for moving from Level 1 to Level 2.

                    Output Requirements:
                    - Return only the recommendation synopsis in plain text without any labels, headers, or introductory 
                    phrases.
                    - Avoid any special formatting, characters, symbols, or Markdown, and do not use labels such as 
                    "Recommendation:" or "Synopsis:".
                    - Do not include any asterisks (*) or stylistic indicators.

                    Sample output:
                    To progress capability maturity, prioritize establishing consistent governance practices, 
                    integrate cloud-specific incident management, and strengthen compliance protocols to align cloud 
                    and overall problem management frameworks.
                """
        # Get the recommendation analysis
        recommendation = ia_analysis(criteria, assessment, prompt)
        recommendation = clean_and_normalize_text(recommendation)
        recommendation = to_sentence_case(recommendation)

        # Construct the backlog prompt
        # prompt = f"""
        #    Build a plain-text numbered list, without any heading or preamble, of 3 implementation activities
        #    that will provide a roadmap between the current state {assessment} to implement the following defined
        #    recommendations: {recommendation}.
        #    The numbered list should be in priority order based on what should be implemented first.
        #    Please provide the output in plain text format without any special headings, characters, symbols, or Markdown-style
        #    formatting. Do not use asterisks (*) for bold text or any other special formatting.
        #    Output Format:
        #    1. Activity 1
        #    2. Activity 2
        #    3. Activity 3
        # """
        # backlog = analysis_modules.build_backlog(prompt, recommendation)
        # backlog = analysis_modules.clean_and_normalize_text(backlog)

        # Save data to the AnalysisResults table in the database
        save_analysis_result(
            project_id=int(project_id),
            capability_id=int(capability_id),
            level=int(level),
            alignment=alignment,
            similarity_score=int(similarity),
            maturity_score=int(maturity_score),
            recommendations=recommendation
            # backlog=backlog
        )

        # Document Output - Commented Out
        # Add Capability heading before analysis
        # heading_a = doc.add_heading(f"{capability}", level=2)
        # if heading_a.runs:
        #     heading_run = heading_a.runs[0]
        #     heading_run.font.size = Pt(14)  # Set heading font size to 16pt

        # Add alignment result to the Word document
        # alignment_paragraph = doc.add_paragraph(f"Alignment result for {capability} (Level {level}): {alignment}")
        # if alignment_paragraph.runs:
        #     alignment_run = alignment_paragraph.runs[0]
        #     alignment_run.font.size = Pt(10)  # Set alignment result font size to 10pt

        # Add similarity score to the Word document
        # similarity_paragraph = doc.add_paragraph(f"Similarity score for {capability} capability: {similarity}")
        # if similarity_paragraph.runs:
        #     similarity_run = similarity_paragraph.runs[0]
        #     similarity_run.font.size = Pt(10)  # Set similarity score font size to 12pt

        # Add maturity score to the Word document
        # maturity_paragraph = doc.add_paragraph(f"Maturity Score for {capability} capability: {maturity_score}")
        # if maturity_paragraph.runs:
        #     maturity_run = maturity_paragraph.runs[0]
        #     maturity_run.font.size = Pt(10)  # Set maturity score font size to 12pt

        # Add recommendations heading to the Word document
        # heading_r = doc.add_heading(f"Recommendations for {capability} capability", level=3)
        # if heading_r.runs:
        #     heading_r_run = heading_r.runs[0]
        #     heading_r_run.font.size = Pt(14)  # Set recommendations heading font size to 14pt

        # Add the recommendation paragraph
        # recommendation_paragraph = doc.add_paragraph(recommendation)
        # if recommendation_paragraph.runs:
        #     recommendation_run = recommendation_paragraph.runs[0]
        #     recommendation_run.font.size = Pt(10)  # Set recommendation font size to 10pt

        # Add backlog heading to the Word document
        # heading_b = doc.add_heading(f"Implementation Road for {capability}", level=3)
        # if heading_b.runs:
        #     heading_b_run = heading_b.runs[0]
        #     heading_b_run.font.size = Pt(14)  # Set the backlog heading font size to 14pt

        # Add backlog as a paragraph in the Word document
        # backlog_paragraph = doc.add_paragraph(backlog)
        # if backlog_paragraph.runs:
        #     backlog_run = backlog_paragraph.runs[0]
        #     backlog_run.font.size = Pt(10)  # Set the backlog text font size to 10pt

        # doc.add_page_break()

    # Document Output - Commented Out
    # Define the file path for the Word document
    # output_file_path = os.path.join(os.getcwd(), f"output/{project_name}_results.docx")

    # Save the Word document once, after all processing
    # doc.save(output_file_path)
    # print(f"Results written to {output_file_path}")
else:
    print(f"No data found for project ID {project_id}.")
